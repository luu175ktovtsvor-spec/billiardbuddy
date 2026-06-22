"""画板成品「落地」：把定稿内容渲染成各格式字节，或写回本地文件。

扩展方式（关键）：要新增一种导出格式，只需写一个 `content:str -> bytes` 函数，
注册进 `_RENDERERS` 即可——其它地方（接口/前端/保存逻辑）都不用改。
今天先支持 md / txt / html / docx；以后加 pptx、pdf 等同理往注册表加一条。
"""
from __future__ import annotations

import io
import re
from html import escape as _esc
from pathlib import Path
from typing import Callable

# 当前支持的导出格式（前端下拉据此；加格式时这里 + _RENDERERS 各加一条）
SUPPORTED: tuple[str, ...] = ("md", "txt", "html", "docx")

# 给前端/对话框用的「格式 → 中文名」
FORMAT_LABELS: dict[str, str] = {"md": "Markdown", "txt": "纯文本", "html": "网页", "docx": "Word 文档"}


def _render_md(content: str) -> bytes:
    return content.encode("utf-8")


def _render_txt(content: str) -> bytes:
    """去掉常见 Markdown 记号，给干净纯文本。"""
    out: list[str] = []
    for ln in content.split("\n"):
        s = re.sub(r"^#{1,6}\s+", "", ln)          # 标题号
        s = re.sub(r"^\s*[-*]\s+", "• ", s)         # 列表点
        s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)      # 粗体
        s = re.sub(r"(?<!\*)\*(?!\*)(.+?)\*", r"\1", s)  # 斜体
        out.append(s)
    return "\n".join(out).encode("utf-8")


def _md_inline_to_html(s: str) -> str:
    """行内：先转义，再还原 **粗** / *斜*（避免 XSS）。"""
    s = _esc(s)
    s = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", s)
    s = re.sub(r"(?<!\*)\*(?!\*)(.+?)\*", r"<em>\1</em>", s)
    return s


def _render_html(content: str) -> bytes:
    """极简 Markdown → 可直接打开的 HTML（标题/段落/列表/粗斜）。"""
    body: list[str] = []
    in_list = False
    for raw in content.split("\n"):
        line = raw.rstrip()
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        is_li = bool(re.match(r"^\s*[-*]\s+", line))
        if not is_li and in_list:
            body.append("</ul>")
            in_list = False
        if m:
            lv = len(m.group(1))
            body.append(f"<h{lv}>{_md_inline_to_html(m.group(2))}</h{lv}>")
        elif is_li:
            if not in_list:
                body.append("<ul>")
                in_list = True
            body.append(f"<li>{_md_inline_to_html(re.sub(r'^\s*[-*]\s+', '', line))}</li>")
        elif line.strip():
            body.append(f"<p>{_md_inline_to_html(line)}</p>")
    if in_list:
        body.append("</ul>")
    html = (
        '<!doctype html><html lang="zh-CN"><head><meta charset="utf-8">'
        '<meta name="viewport" content="width=device-width,initial-scale=1">'
        "<style>body{font-family:-apple-system,system-ui,'PingFang SC',sans-serif;"
        "max-width:720px;margin:40px auto;padding:0 20px;line-height:1.75;color:#1d1d1f;}"
        "h1{font-size:24px;}h2{font-size:20px;}h3{font-size:17px;}</style></head><body>\n"
        + "\n".join(body)
        + "\n</body></html>"
    )
    return html.encode("utf-8")


def _render_docx(content: str) -> bytes:
    """Markdown → .docx（标题/列表/段落）。富文本保真够日常办公用。"""
    import docx

    d = docx.Document()
    for raw in content.split("\n"):
        line = raw.rstrip()
        if not line.strip():
            continue
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            d.add_heading(m.group(2).strip(), level=min(len(m.group(1)), 4))
        elif re.match(r"^\s*[-*]\s+", line):
            d.add_paragraph(re.sub(r"^\s*[-*]\s+", "", line).strip(), style="List Bullet")
        else:
            # 去掉 **/*，正文按普通段落（保真足够；要更细可后续按 run 拆粗体）
            d.add_paragraph(re.sub(r"\*{1,2}(.+?)\*{1,2}", r"\1", line))
    bio = io.BytesIO()
    d.save(bio)
    return bio.getvalue()


# 注册表：格式 → 渲染函数。新增格式只动这里 + SUPPORTED。
_RENDERERS: dict[str, Callable[[str], bytes]] = {
    "md": _render_md,
    "txt": _render_txt,
    "html": _render_html,
    "docx": _render_docx,
}


def render_bytes(content: str, fmt: str) -> bytes:
    """把成品内容渲染成指定格式的字节。"""
    fn = _RENDERERS.get(fmt)
    if fn is None:
        raise ValueError(f"不支持的导出格式：{fmt}")
    return fn(content)


def _safe_name(name: str) -> str:
    """文件名清洗：去路径分隔/非法字符，避免越界。"""
    base = (name or "").strip().split("\n")[0]
    base = re.sub(r'[\\/:*?"<>|]', "", base).strip()
    return (base or "成品")[:60]


def save_to_library(content: str, fmt: str, name: str) -> Path:
    """存进「内容库/成品」，重名先备份。返回写入路径。"""
    from services.agent.local_tools import _backup, _library_root

    if fmt not in _RENDERERS:
        raise ValueError(f"不支持的导出格式：{fmt}")
    root = _library_root() / "成品"
    root.mkdir(parents=True, exist_ok=True)
    path = root / f"{_safe_name(name)}.{fmt}"
    if path.exists():
        _backup(path)
    path.write_bytes(render_bytes(content, fmt))
    return path
